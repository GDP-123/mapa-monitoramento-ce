import folium, json, sqlite3, base64
import pandas as pd
import streamlit as st

from shapely.geometry import shape, Point
from collections import defaultdict
from folium import LayerControl
from folium.features import GeoJsonTooltip
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Função de filtro segura
def contem_cidade(cidades_lista, cidade_alvo):
    if not isinstance(cidades_lista, list):  # Caso algum valor não seja lista
        return False
    return any(cidade_alvo.lower() == c.lower() for c in cidades_lista)

# EM DESUSO NO MOMENTO
def camada_colorida(dados, m, ce_geo):
    # Cria o dicionário cidade -> lista de facções
    cidade_faccao = defaultdict(list)
    for pessoa in dados:
        # Agora pessoa['cidade'] é uma lista, então iteramos sobre cada cidade
        for cidade in pessoa['cidade']:
            faccao = pessoa['faccao']
            cidade_faccao[cidade].append(faccao)

    # Opcional: converter para dict comum se não quiser usar defaultdict
    cidade_faccao = dict(cidade_faccao)

    # Determinar a facção predominante em cada cidade
    faccao_predominante = {}
    for cidade, faccoes in cidade_faccao.items():
        # Contar ocorrências de cada facção
        contagem = {}
        for f in faccoes:
            contagem[f] = contagem.get(f, 0) + 1
        # Pegar a facção mais comum
        faccao_predominante[cidade] = max(contagem.items(), key=lambda x: x[1])[0]

    # Cores para cada facção (mantive suas cores originais)
    cores_faccoes = {
        "Facção A": "red",
        "Facção B": "blue",
        "Facção C": "green",
        "Facção D": "purple",
        "Facção E": "orange"
    }

    # Adicionar cada facção como uma camada separada
    for faccao, cor in cores_faccoes.items():
        # Filtrar cidades desta facção
        cidades_faccao = [
            feature for feature in ce_geo['features']
            if feature['properties']['name'] in faccao_predominante and 
            faccao_predominante[feature['properties']['name']] == faccao
        ]
        
        if not cidades_faccao:
            continue
        
        # Criar GeoJSON para esta facção
        geojson_faccao = {
            "type": "FeatureCollection",
            "features": cidades_faccao
        }
        
        # Adicionar ao mapa como camada
        folium.GeoJson(
            geojson_faccao,
            name=faccao,
            style_function=lambda x, cor=cor: {
                'fillColor': cor,
                'color': 'black',
                'weight': 0.5,
                'fillOpacity': 0.6
            },
            tooltip=GeoJsonTooltip(
                fields=['name'], 
                aliases=['Cidade:'],
                labels=True
            )
        ).add_to(m)

    # Adicionar controle de camadas
    folium.LayerControl().add_to(m)

    return m

# Contorno da região
def camada_amarela(ce_geo,m):
    geojson_layer = folium.GeoJson(
        ce_geo,
        name="Cidades",
        tooltip=folium.GeoJsonTooltip(fields=['name'], aliases=['Cidade:']),
        style_function=lambda x: {
            'fillColor': '#ffffff',
            'color': 'black',
            'weight': 0.5,
            'fillOpacity': 0.1,
        },
        highlight_function=lambda x: {
            'weight': 3,
            'color': 'yellow'
        },  # quando passar o mouse
    )
    geojson_layer.add_to(m)
    return geojson_layer

# Desenhar polígonos
def poligonos_coloridos(cores_faccoes, dados, m, faccoes_selecionadas):
    # Criar grupos apenas para as facções selecionadas
    grupos = {}
    for faccao in faccoes_selecionadas:  # Alteração aqui: percorre apenas as selecionadas
        if faccao in cores_faccoes:  # Verifica se a facção existe no dicionário de cores
            grupos[faccao] = folium.FeatureGroup(name=faccao)

    # Adicionar cada indivíduo ao grupo correspondente (se a facção estiver selecionada)
    for individuo in dados:
        faccao_individuo = individuo.get("faccao")
        
        # Verifica se a facção do indivíduo está selecionada e existe nos grupos
        if faccao_individuo in grupos:
            if "geo" in individuo and individuo["geo"]["tipo"] == "circle":
                folium.Circle(
                    location=[individuo["geo"]["lat"], individuo["geo"]["long"]],
                    radius=individuo["geo"]["radius"] * 1,
                    color=cores_faccoes.get(faccao_individuo, "gray"),
                    fill=True,
                    fill_opacity=0.4,
                    tooltip=f"{individuo['nome']} ({faccao_individuo})"
                ).add_to(grupos[faccao_individuo])

            elif "geo" in individuo and individuo["geo"]["tipo"] == "polygon" and "pontos" in individuo["geo"]:
                folium.Polygon(
                    locations=individuo["geo"]["pontos"],
                    color=cores_faccoes.get(faccao_individuo, "gray"),
                    fill=True,
                    fill_opacity=0.2,
                    weight=2,
                    tooltip=f"{individuo['nome']} ({faccao_individuo})"
                ).add_to(grupos[faccao_individuo])

    # Adicionar apenas os grupos das facções selecionadas ao mapa
    for faccao, grupo in grupos.items():
        grupo.add_to(m)

    return m

# Função para encontrar a cidade pelo ponto
def encontrar_cidade_por_coordenada(lat, lng, geojson):
    ponto = Point(lng, lat)  # Atenção: Point usa (x=lng, y=lat)
    
    for feature in geojson["features"]:
        poligono = shape(feature["geometry"])
        if poligono.contains(ponto):
            return feature["properties"]["name"]  # ou outro campo que quiser
    return None

# Função para mostrar/ocultar a mensagem
def toggle_mensagem():
    if 'mostrar_mensagem' not in st.session_state:
        st.session_state.mostrar_mensagem = False
    st.session_state.mostrar_mensagem = not st.session_state.mostrar_mensagem

#Criar relatório PDF
def gerar_documento(modelo_path, regiao, dados_tabela):
    try:             
        # Abre o documento modelo
        document = Document(modelo_path)
        
        # Adiciona título (usando style diretamente para evitar problemas)
        document.add_heading(f'Relatório de Inteligência - {regiao}/CE', 0)
        
        # Adiciona data atual formatada
        data_atual = datetime.now().strftime('%d/%m/%Y')
        data_paragraph = document.add_paragraph(data_atual)
        data_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Adiciona texto justificado
        texto = document.add_paragraph()
         # Adiciona título (nível 0 é o mais alto)
        texto.add_run(f'A seguir são apresentados dados vinculados às lideranças de organizações criminosas do município de {regiao}/CE.')
        texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        #Criando a tabela   
        tabela = document.add_table(rows=1, cols=6)
        tabela.alignment = WD_TABLE_ALIGNMENT.LEFT
        tabela.autofit = False
        
        # Configura largura das colunas
        widths = (Inches(0.8), Inches(1.5), Inches(1.2), Inches(1.2), Inches(1), Inches(1.2))
        for i, width in enumerate(widths):
            tabela.columns[i].width = width
        
        # Cabeçalho
        cabecalho = tabela.rows[0].cells
        cabecalho_campos = ['Foto', 'Nome', 'Facção', 'Cargo', 'Status', 'Quantidade Mandatos']
        for i, campo in enumerate(cabecalho_campos):
            cabecalho[i].text = campo
            # Formata cabeçalho
            for paragraph in cabecalho[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Adiciona os dados
        for dados in dados_tabela.itertuples(index=False):

            linha = tabela.add_row().cells
            linha[1].text = dados.nome
            linha[2].text = dados.faccao
            linha[3].text = dados.cargo
            linha[4].text = dados.status
            linha[5].text = str(dados.quantidade_mandados)
        
        # Adiciona seção de atualizações com formatação robusta
        atualizacoes_title = document.add_paragraph()
        atualizacoes_title_run = atualizacoes_title.add_run('\nAtualizações')
        atualizacoes_title_run.bold = True
        atualizacoes_title_run.font.size = Pt(14)
        
        # Adiciona atualização com formatação
        atualizacao = document.add_paragraph()
        
        # Data em negrito e cor diferente
        run_data = atualizacao.add_run('02/05/2025 — ')
        run_data.bold = True
        run_data.font.color.rgb = RGBColor(0x42, 0x51, 0x59)  # Azul
        
        # Texto normal
        run_texto = atualizacao.add_run('Liderança do grupo comando vermelho da região de sobral foi presa. '
                                      'Acredita-se que isso irá gerar uma onda de retaliações na cidade. '
                                      'Fonte: DIP/PCCE')
        
        # Alinha justificado
        atualizacao.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        #Salvando arquivo docx
        #document.save(output_path)
        return document

    except Exception as e:
        print(f"Erro ao gerar documento: {str(e)}")
        raise
    
# Função para criar link de download
def get_download_link(pdf_output, filename):
    b64 = base64.b64encode(pdf_output).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download PDF</a>'
    return href

# Carregar na memoria
@st.cache_data
def load_data():
    with open('dados.json', 'r', encoding='utf-8') as f:
        dados = json.load(f)
    df = pd.DataFrame(dados)

    # Carregar GeoJSON
    with open('ce_cities.geojson', 'r', encoding='utf-8') as f:
        ce_geo = json.load(f)

    return df,dados,ce_geo

# Inicializa o banco de dados
def init_db():
    conn = sqlite3.connect('notificacoes.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS notificacoes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            mensagem TEXT NOT NULL,
            data TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

# Adiciona uma notificação
# Exemplo: inserir notificações manualmente para testes
#with st.expander("Adicionar nova notificação"):
#    nova_msg = st.text_area("Mensagem da notificação")
#    if st.button("Adicionar"):
#        if nova_msg.strip():
#            adicionar_notificacao(nova_msg.strip())
#            st.success("Notificação adicionada!")
#            st.experimental_user()
#        else:
#            st.error("A mensagem não pode estar vazia.")
def adicionar_notificacao(mensagem):
    conn = sqlite3.connect('notificacoes.db')
    c = conn.cursor()
    data_atual = datetime.now().strftime("%d/%m/%Y")
    c.execute('INSERT INTO notificacoes (mensagem, data) VALUES (?, ?)', (mensagem, data_atual))
    conn.commit()
    conn.close()

# Buscar notificacao
def buscar_notificacoes(limit=10):
    conn = sqlite3.connect('notificacoes.db')
    c = conn.cursor()
    c.execute('SELECT mensagem, data FROM notificacoes ORDER BY id DESC LIMIT ?', (limit,))
    rows = c.fetchall()
    conn.close()
    return rows
